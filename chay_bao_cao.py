import os
import sys
import subprocess

# Reconfigure stdout to support UTF-8 Vietnamese output in Windows terminal
if sys.platform.startswith('win'):
    try:
        sys.stdout.reconfigure(encoding='utf-8')
        sys.stderr.reconfigure(encoding='utf-8')
    except AttributeError:
        pass

# Auto-install python-docx if not installed
try:
    import docx
except ImportError:
    print("Đang cài đặt thư viện python-docx...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_page_number(run):
    """Inserts a page number field inside a run."""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def style_run(run, font_name='Times New Roman', font_size=13, bold=False, italic=False, color_rgb=None):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    if color_rgb:
        run.font.color.rgb = color_rgb

def add_formatted_paragraph(doc, text="", style=None, font_size=13, bold=False, italic=False, color_rgb=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph(style=style)
    p.alignment = align
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(0)
    
    # Parse bold formatting **bold**
    if "**" in text:
        parts = text.split("**")
        for idx, part in enumerate(parts):
            if idx % 2 == 1:
                run = p.add_run(part)
                style_run(run, font_size=font_size, bold=True, italic=italic, color_rgb=color_rgb)
            else:
                run = p.add_run(part)
                style_run(run, font_size=font_size, bold=bold, italic=italic, color_rgb=color_rgb)
    else:
        run = p.add_run(text)
        style_run(run, font_size=font_size, bold=bold, italic=italic, color_rgb=color_rgb)
    return p

def create_cover_page(doc):
    """Creates a professional cover page for the project report."""
    # School Header
    p1 = add_formatted_paragraph(doc, "BỘ GIÁO DỤC VÀ ĐÀO TẠO", font_size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    p2 = add_formatted_paragraph(doc, "TRƯỜNG ĐẠI HỌC CÔNG NGHỆ", font_size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    p3 = add_formatted_paragraph(doc, "KHOA CÔNG NGHỆ THÔNG TIN", font_size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    
    # Add separator line
    p_sep = doc.add_paragraph()
    p_sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sep.paragraph_format.space_before = Pt(12)
    p_sep.paragraph_format.space_after = Pt(24)
    run_sep = p_sep.add_run("-------------------- *** --------------------")
    style_run(run_sep, font_size=12, color_rgb=RGBColor(148, 163, 184))
    
    # Add large vertical spacing
    for _ in range(3):
        doc.add_paragraph()
        
    # Project Title
    p_subj = add_formatted_paragraph(doc, "ĐỒ ÁN MÔN HỌC LẬP TRÌNH PYTHON", font_size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    p_subj.paragraph_format.space_after = Pt(12)
    
    p_title = add_formatted_paragraph(
        doc, 
        "ĐỀ TÀI: NGHIÊN CỨU NGÔN NGỮ PYTHON VÀ LẬP TRÌNH GAME TRƯỢT HÌNH \"PUZZLE GAME\" BẰNG THƯ VIỆN PYGAME", 
        font_size=18, 
        bold=True, 
        color_rgb=RGBColor(13, 148, 136), # Teal
        align=WD_ALIGN_PARAGRAPH.CENTER
    )
    p_title.paragraph_format.space_after = Pt(48)
    
    # Vertical spacing
    for _ in range(4):
        doc.add_paragraph()
        
    # Student Info
    info_table = doc.add_table(rows=4, cols=2)
    info_table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    
    info_data = [
        ("Giảng viên hướng dẫn:", "Thầy Cô Khoa CNTT"),
        ("Sinh viên thực hiện:", "Đỗ Xuân Trung"),
        ("Mã số sinh viên:", "3025110405"),
        ("Lớp / Khóa học:", "TH30.06 / Đồ án phần mềm Python")
    ]
    
    for row_idx, (label, val) in enumerate(info_data):
        row = info_table.rows[row_idx]
        
        # Label cell
        cell_lbl = row.cells[0]
        cell_lbl.width = Inches(2.5)
        p_lbl = cell_lbl.paragraphs[0]
        p_lbl.paragraph_format.line_spacing = 1.3
        run_lbl = p_lbl.add_run(label)
        style_run(run_lbl, font_size=13, bold=(row_idx == 1))
        
        # Value cell
        cell_val = row.cells[1]
        cell_val.width = Inches(3.5)
        p_val = cell_val.paragraphs[0]
        p_val.paragraph_format.line_spacing = 1.3
        run_val = p_val.add_run(val)
        style_run(run_val, font_size=13, bold=(row_idx == 1))
        
    # Vertical spacing
    for _ in range(4):
        doc.add_paragraph()
        
    # Footer Location/Date
    p_foot = add_formatted_paragraph(doc, "Thành phố Hồ Chí Minh, Năm 2026", font_size=13, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    
    # Page Break
    doc.add_page_break()

def markdown_to_docx(md_path, docx_path):
    doc = Document()
    
    # Clear headers for all sections to prevent any text/page numbering in header
    for sec in doc.sections:
        sec.header.is_linked_to_previous = False
        for p in sec.header.paragraphs:
            p.text = ""
            p.clear()
            
    # Page setup - A4 paper size
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    
    # Margins: Top 2.0cm, Bottom 2.0cm, Left 3.0cm, Right 2.0cm (Exactly using Cm unit)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)
    
    # Enable different first page to hide header/footer on cover
    section.different_first_page_header_footer = True
    
    # Add page number in footer for standard pages
    footer = section.footer
    f_para = footer.paragraphs[0]
    f_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f_run = f_para.add_run()
    style_run(f_run, font_name='Times New Roman', font_size=11)
    add_page_number(f_run)
    
    # Create cover page
    create_cover_page(doc)
    
    print(f"Đọc dữ liệu từ {md_path}...")
    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()
        
    in_code_block = False
    in_table = False
    table_rows = []
    
    for line in lines:
        stripped = line.strip()
        
        # Skip empty lines unless inside code block
        if not stripped:
            if in_code_block:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.space_before = Pt(0)
            continue
            
        # Skip plain markdown separators
        if stripped == "---":
            continue
            
        # Handle code blocks (No Courier New! Use Times New Roman 11pt, italicized and indented)
        if stripped.startswith("```"):
            in_code_block = not in_code_block
            continue
            
        if in_code_block:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(0)
            run = p.add_run(line.rstrip('\n'))
            style_run(run, font_name='Times New Roman', font_size=11, italic=True, color_rgb=RGBColor(100, 116, 139))
            continue
            
        # Handle tables (like the abbreviations table)
        if stripped.startswith("|") and stripped.endswith("|"):
            in_table = True
            # Skip separator lines e.g. | :--- | :--- |
            if "---" in stripped:
                continue
            # Parse table cells
            cells = [c.strip() for c in stripped.split("|")[1:-1]]
            table_rows.append(cells)
            continue
        elif in_table:
            # Table ended, render it!
            in_table = False
            if table_rows:
                num_cols = len(table_rows[0])
                table = doc.add_table(rows=len(table_rows), cols=num_cols)
                table.style = 'Table Grid'
                table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
                
                for r_idx, row_data in enumerate(table_rows):
                    row = table.rows[r_idx]
                    for c_idx, cell_value in enumerate(row_data):
                        cell = row.cells[c_idx]
                        p = cell.paragraphs[0]
                        p.paragraph_format.space_after = Pt(4)
                        p.paragraph_format.space_before = Pt(4)
                        p.paragraph_format.line_spacing = 1.15
                        run = p.add_run(cell_value)
                        # Header bold, body normal
                        style_run(run, font_name='Times New Roman', font_size=12, bold=(r_idx == 0))
                
                # Add spacing after table
                p_space = doc.add_paragraph()
                p_space.paragraph_format.space_before = Pt(6)
                p_space.paragraph_format.space_after = Pt(6)
                table_rows = []
                
        # Handle image markup ![Caption](path)
        if stripped.startswith("![") and "](" in stripped and stripped.endswith(")"):
            caption_start = stripped.find("[") + 1
            caption_end = stripped.find("](")
            path_start = caption_end + 2
            path_end = len(stripped) - 1
            
            caption = stripped[caption_start:caption_end]
            img_path = stripped[path_start:path_end]
            
            if os.path.exists(img_path):
                print(f"Đang chèn hình ảnh: {img_path}...")
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.paragraph_format.space_before = Pt(12)
                p_img.paragraph_format.space_after = Pt(6)
                
                # Insert picture with width fitting nicely on A4 page (5.0 inches width = 12.7 cm)
                run_img = p_img.add_run()
                run_img.add_picture(img_path, width=Inches(5.0))
                
                # Insert centered caption below image
                p_cap = doc.add_paragraph()
                p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_cap.paragraph_format.space_before = Pt(4)
                p_cap.paragraph_format.space_after = Pt(12)
                run_cap = p_cap.add_run(caption)
                style_run(run_cap, font_name='Times New Roman', font_size=11, italic=True)
            else:
                print(f"[CẢNH BÁO] Không tìm thấy hình ảnh tại: {img_path}")
            continue
            
        # Headings (All Times New Roman, bold, size according to level)
        if stripped.startswith("# "):
            h_text = stripped[2:]
            if h_text == "BÁO CÁO ĐỒ ÁN MÔN HỌC":
                continue # Skip duplicate title
            
            # Start each main chapter on a new page
            doc.add_page_break()
            
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(24)
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(h_text)
            style_run(run, font_name='Times New Roman', font_size=20, bold=True, color_rgb=RGBColor(13, 148, 136))
        elif stripped.startswith("## "):
            h_text = stripped[3:]
            if h_text.startswith("ĐỀ TÀI:"):
                continue # Skip duplicate subtitle
                
            if h_text == "---":
                continue # Skip markdown line separator
                
            # Trigger page break for specific front-matter and back-matter sections
            sections_to_break = ["MỤC LỤC", "DANH MỤC TỪ VIẾT TẮT", "DANH MỤC HÌNH ẢNH", "MỞ ĐẦU", "KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN", "TÀI LIỆU THAM KHẢO"]
            if h_text in sections_to_break:
                doc.add_page_break()
                
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(h_text)
            style_run(run, font_name='Times New Roman', font_size=15, bold=True, color_rgb=RGBColor(71, 85, 105))
        elif stripped.startswith("### "):
            h_text = stripped[4:]
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(h_text)
            style_run(run, font_name='Times New Roman', font_size=13, bold=True)
        elif stripped.startswith("#### "):
            h_text = stripped[5:]
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(h_text)
            style_run(run, font_name='Times New Roman', font_size=13, bold=True, italic=True)
            
        # Bullet Lists
        elif stripped.startswith("* ") or stripped.startswith("- "):
            list_text = stripped[2:]
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.line_spacing = 1.3
            
            if "**" in list_text:
                parts = list_text.split("**")
                for idx, part in enumerate(parts):
                    if idx % 2 == 1:
                        run = p.add_run(part)
                        style_run(run, font_name='Times New Roman', font_size=13, bold=True)
                    else:
                        run = p.add_run(part)
                        style_run(run, font_name='Times New Roman', font_size=13)
            else:
                run = p.add_run(list_text)
                style_run(run, font_name='Times New Roman', font_size=13)
                
        # Normal paragraphs (Always Times New Roman, size 13, line spacing 1.5)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.line_spacing = 1.5
            
            if "**" in stripped:
                parts = stripped.split("**")
                for idx, part in enumerate(parts):
                    if idx % 2 == 1:
                        run = p.add_run(part)
                        style_run(run, font_name='Times New Roman', font_size=13, bold=True)
                    else:
                        run = p.add_run(part)
                        style_run(run, font_name='Times New Roman', font_size=13)
            else:
                run = p.add_run(stripped)
                style_run(run, font_name='Times New Roman', font_size=13)
                
    print(f"Đang lưu file Word tại: {docx_path}...")
    doc.save(docx_path)
    print("Xuất báo cáo thành công!")

if __name__ == "__main__":
    BASE_DIR = os.path.dirname(os.path.abspath(__file__))
    md_file = os.path.join(BASE_DIR, "Bao_Cao_Do_An.md")
    
    # Cập nhật tệp Word báo cáo duy nhất
    filename = "Báo Cáo Thiết Kế.docx"
    dest_path = os.path.join(BASE_DIR, filename)
    try:
        print(f"Đang ghi đè lên file: {filename}...")
        markdown_to_docx(md_file, dest_path)
        print(f"-> Xuất báo cáo thành công cho: {filename}\n")
    except PermissionError:
        print(f"\n[CẢNH BÁO] Không thể ghi đè lên file '{filename}' vì file đang mở trong Microsoft Word.")
        print("Vui lòng ĐÓNG Microsoft Word lại và chạy lại file này để cập nhật báo cáo mới nhất!\n")
